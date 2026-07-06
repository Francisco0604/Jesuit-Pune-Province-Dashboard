import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = create_element('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = create_element('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = create_element(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_border(cell, **kwargs):
    """
    kwargs can be top, bottom, left, right.
    values are dicts: {'sz': 12, 'val': 'single', 'color': 'FF0000', 'space': '0'}
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = create_element('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f'w:{edge}'
            element = create_element(tag)
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            element.set(qn('w:space'), str(edge_data.get('space', 0)))
            element.set(qn('w:color'), edge_data.get('color', 'auto'))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def make_clear_cell(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = create_element('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = f'w:{edge}'
        element = create_element(tag)
        element.set(qn('w:val'), 'none')
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_horizontal_border(paragraph, color_hex="C5A059", sz="12", space="4"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = create_element('w:pBdr')
    bottom = create_element('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:space'), space)
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def make_flowchart_node(cell, title, description, border_color="A34436", bg_color="F9F8F6"):
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    set_cell_border(cell,
        top={'sz': 12, 'val': 'single', 'color': border_color},
        bottom={'sz': 12, 'val': 'single', 'color': border_color},
        left={'sz': 12, 'val': 'single', 'color': border_color},
        right={'sz': 12, 'val': 'single', 'color': border_color}
    )
    
    # Title
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run_title = p.add_run(title)
    run_title.font.name = 'Georgia'
    run_title.font.size = Pt(12)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(163, 68, 54) # Terracotta
    
    # Description
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.1
    run_desc = p2.add_run(description)
    run_desc.font.name = 'Calibri'
    run_desc.font.size = Pt(10)
    run_desc.font.color.rgb = RGBColor(45, 45, 45) # Charcoal

def make_arrow_cell(cell, arrow_char="↓"):
    make_clear_cell(cell)
    set_cell_margins(cell, top=40, bottom=40, left=40, right=40)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(arrow_char)
    run.font.name = 'Georgia'
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = RGBColor(197, 160, 89) # Gold

def build_flowchart_docx():
    doc = Document()
    
    # Set Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Document Header Page Elements
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(15)
    title_p.paragraph_format.space_after = Pt(6)
    title_run = title_p.add_run("JESUIT PUNE PROVINCE ACTIVITY MAP")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.name = 'Georgia'
    title_run.font.color.rgb = RGBColor(163, 68, 54) # Terracotta
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(18)
    sub_run = sub_p.add_run("Website Layout & Visitor Navigation Flowchart")
    sub_run.italic = True
    sub_run.font.size = Pt(13)
    sub_run.font.name = 'Georgia'
    sub_run.font.color.rgb = RGBColor(128, 128, 128)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Thick border separating header
    border_p = doc.add_paragraph()
    border_p.paragraph_format.space_after = Pt(20)
    add_horizontal_border(border_p, color_hex="A34436", sz="18", space="1")
    
    # Intro
    intro_p = doc.add_paragraph()
    intro_p.paragraph_format.space_after = Pt(24)
    intro_run = intro_p.add_run(
        "Below is the visual flowchart of the website. It illustrates the layout of each page "
        "and the navigation path a visitor follows when exploring the map."
    )
    intro_run.font.name = 'Calibri'
    intro_run.font.size = Pt(11)
    intro_run.font.color.rgb = RGBColor(45, 45, 45)
    
    # Create the Flowchart Table (3 Columns)
    # Col 0: Left Branch (3.0 in), Col 1: Center Spacer (0.8 in), Col 2: Right Branch (3.0 in)
    table = doc.add_table(rows=10, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    col_widths = [Inches(3.0), Inches(0.8), Inches(3.0)]
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width
            
    # Row 0: Landing Page (Merged)
    cell_r0 = table.rows[0].cells[0].merge(table.rows[0].cells[2])
    make_flowchart_node(
        cell_r0,
        "1. Home / Landing Page (/) ",
        "The visitor entry point. Introduces the Jesuit mission, AMDG motto, presence stats, and links to the interactive map."
    )
    
    # Row 1: Arrow Down
    cell_r1 = table.rows[1].cells[0].merge(table.rows[1].cells[2])
    make_arrow_cell(cell_r1, "↓")
    
    # Row 2: Interactive Map Page (Merged)
    cell_r2 = table.rows[2].cells[0].merge(table.rows[2].cells[2])
    make_flowchart_node(
        cell_r2,
        "2. Interactive Map Page (/map)",
        "The core display. Features a Leaflet map showing color-coded markers and boundaries, with a sidebar for search and filtering."
    )
    
    # Row 3: Split Arrows
    make_arrow_cell(table.rows[3].cells[0], "↙")
    make_clear_cell(table.rows[3].cells[1])
    make_arrow_cell(table.rows[3].cells[2], "↘")
    
    # Row 4: Path Headers
    make_flowchart_node(
        table.rows[4].cells[0],
        "VISITOR PATH\n(Explore & Search)",
        "For general public users, scholars, and parish members visiting the site.",
        border_color="C5A059" # Gold border
    )
    make_clear_cell(table.rows[4].cells[1])
    make_flowchart_node(
        table.rows[4].cells[2],
        "ADMINISTRATOR PATH\n(Manage & Update)",
        "For authorized coordinators and web developers updating province records.",
        border_color="C5A059" # Gold border
    )
    
    # Row 5: Arrows Down
    make_arrow_cell(table.rows[5].cells[0], "↓")
    make_clear_cell(table.rows[5].cells[1])
    make_arrow_cell(table.rows[5].cells[2], "↓")
    
    # Row 6: Actions
    make_flowchart_node(
        table.rows[6].cells[0],
        "3A. Search & Click Markers",
        "Visitor types a village name or filters categories (like parishes or NFE centres) and clicks a map pin."
    )
    make_clear_cell(table.rows[6].cells[1])
    make_flowchart_node(
        table.rows[6].cells[2],
        "3B. Access Admin Panel (/admin)",
        "Staff navigates to the secure Admin Portal dashboard to view statistics and database controls."
    )
    
    # Row 7: Arrows Down
    make_arrow_cell(table.rows[7].cells[0], "↓")
    make_clear_cell(table.rows[7].cells[1])
    make_arrow_cell(table.rows[7].cells[2], "↓")
    
    # Row 8: Final Screens
    make_flowchart_node(
        table.rows[8].cells[0],
        "4A. View Details Drawer",
        "A panel slides out from the right showing establishing year, family count, total population, and custom statistics."
    )
    make_clear_cell(table.rows[8].cells[1])
    make_flowchart_node(
        table.rows[8].cells[2],
        "4B. Database Management",
        "Coordinators export statistics to Excel files, import updated spreadsheets, or upload spatial maps (GeoJSON)."
    )
    
    # Row 9: Empty row spacer to clean up spacing
    for c in table.rows[9].cells:
        make_clear_cell(c)
        
    doc.add_paragraph() # Spacer
    
    # Save Document
    dest_dir = "C:\\Users\\franc\\Projects\\Map\\Code"
    dest_path = os.path.join(dest_dir, "WEBSITE_FLOWCHART.docx")
    doc.save(dest_path)
    print(f"Successfully generated flowchart docx at: {dest_path}")
    return dest_path

if __name__ == "__main__":
    build_flowchart_docx()
