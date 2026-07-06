import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# Set Margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title
title_p = doc.add_paragraph()
title_run = title_p.add_run("JESUIT PUNE PROVINCE MAP - TASK CHECKLIST")
title_run.bold = True
title_run.font.size = Pt(18)
title_run.font.name = 'Georgia'
title_run.font.color.rgb = RGBColor(163, 68, 54) # Terracotta
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
sub_p = doc.add_paragraph()
sub_run = sub_p.add_run("Implementation Checklists Grouped by Responsibility")
sub_run.italic = True
sub_run.font.size = Pt(11)
sub_run.font.name = 'Calibri'
sub_run.font.color.rgb = RGBColor(128, 128, 128)
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Person Lists
people_tasks = [
    ("1. JAUNITA - INFORMATION CHECKLIST", [
        "Download the pre-populated CSV template from the Admin Panel or locate it at public/data/village_import_template.csv.",
        "Edit and verify the statistics (e.g., Students, Active Teachers, Families) for all 36 NFE Centres and parishes in Microsoft Excel.",
        "Confirm that tehsil and cluster names match the existing records.",
        "Upload and import the completed Excel CSV file through the Admin Dashboard to link statistics and update village profiles."
    ]),
    ("2. JUHI - GIS & SPATIAL DATA CHECKLIST", [
        "Gather GeoJSON point or polygon files for any missing villages or boundaries.",
        "Upload these GeoJSON files through the Admin Dashboard to initialize geographical shapes and set coordinate locations on the map."
    ]),
    ("3. FRANCISCO - DEVELOPMENT & DEPLOYMENT CHECKLIST", [
        "Security & Access Control: Secure the Admin Dashboard (/admin) using NextAuth.js or basic HTTP Auth so that public users cannot write or delete records.",
        "Data Validation & Sanitation: Build server-side validation on CSV file imports to prevent incorrect types, missing names, or non-numeric values.",
        "Automatic Data Backups: Implement an automatic backup script for public/data/Village_details and public/data/uploads directories.",
        "Senior Accessibility & Contrast Audit: Check readability, font sizing, and color contrast ratios for senior members of the province (contrast between #f5f5dc parchment and text).",
        "Mobile Optimization: Verify drawer and filter sidebar responsiveness on mobile devices/tablets.",
        "Offline Caching: Configure Leaflet tile caching to allow map access in low-connectivity zones.",
        "Code Quality & Warning Resolution: Run 'npm run lint' to scan and clean up styling/code warnings, and fix Next.js hydration issues common with Leaflet SSR dynamic rendering.",
        "Production Build & Bundling: Run 'npm run build' to test bundling and static generation.",
        "Hosting & Deployment Setup: Deploy the repository to a production environment (Vercel, Netlify, or VPS).",
        "Custom Domain & SSL: Direct the custom domain name (e.g., map.punejesuits.org) and enable HTTPS/SSL certificate."
    ]),
    ("4. FRANCISCO - MAP BEAUTIFICATION & ADVANCED FEATURES CHECKLIST", [
        "Change Leaflet Map Tile Provider: Swap the default OpenStreetMap tile layer for a premium, clean, and highly aesthetic tile layout (e.g., CartoDB Positron, Stadia Alidade Smooth, or a custom parchment-styled theme) that matches the site design.",
        "Mark District & Inner Boundaries: Load and render GeoJSON boundaries overlays to draw distinct lines marking the district borders and the inner taluka/tehsil boundaries on the map.",
        "Beautify Map Aesthetics & Interactivity: Customize the marker cluster styling to use the province theme colors, replace default Leaflet popups with custom HTML/CSS styled popup cards, and add interactive hover highlights to boundary shapes.",
        "Dynamic News Administration System: Move the news items in /news from static code into dynamic JSON profiles, and add a section in the Admin Dashboard to manage news articles.",
        "Search Fuzzy Matching: Enhance the sidebar search filter in /map with fuzzy matching so that minor typos in village/district names still return search results."
    ])
]

for name, tasks in people_tasks:
    h2_p = doc.add_heading(level=2)
    h2_p.paragraph_format.space_before = Pt(14)
    h2_p.paragraph_format.space_after = Pt(6)
    h2_run = h2_p.add_run(name)
    h2_run.font.size = Pt(12)
    h2_run.font.name = 'Georgia'
    h2_run.font.color.rgb = RGBColor(163, 68, 54) # Terracotta
    
    for task in tasks:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(4)
        
        box_run = p.add_run("[ ] ")
        box_run.bold = True
        box_run.font.name = 'Calibri'
        box_run.font.color.rgb = RGBColor(197, 160, 89) # Gold color
        
        text_run = p.add_run(task)
        text_run.font.name = 'Calibri'

dest_path = os.path.join(os.path.dirname(__file__), "..", "task_list_by_person_v2.docx")
doc.save(dest_path)
print(f"Successfully generated docx at: {dest_path}")
